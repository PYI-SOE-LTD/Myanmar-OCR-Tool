import os
import re
import subprocess
import threading
import tkinter as tk
from pathlib import Path
from tkinter import filedialog, messagebox, ttk


IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}


def natural_key(text: str):
    return [int(part) if part.isdigit() else part.lower() for part in re.split(r"(\d+)", text)]


def list_images(folder: Path):
    files = [p for p in folder.iterdir() if p.is_file() and p.suffix.lower() in IMAGE_EXTS]
    return sorted(files, key=lambda p: natural_key(p.name))


def extract_trailing_number(stem: str):
    match = re.search(r"(\d+)$", stem)
    if not match:
        return None
    return int(match.group(1))


def parse_page_range(raw: str):
    text = (raw or "").strip()
    if not text:
        return None

    ranges = []
    for part in text.split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            left, right = part.split("-", 1)
            left = left.strip()
            right = right.strip()
            start = int(left) if left else None
            end = int(right) if right else None
            ranges.append((start, end))
        else:
            value = int(part)
            ranges.append((value, value))
    return ranges


def number_in_ranges(number: int, ranges):
    if ranges is None:
        return True
    for start, end in ranges:
        if start is None and end is None:
            return True
        if start is None and number <= end:
            return True
        if end is None and number >= start:
            return True
        if start <= number <= end:
            return True
    return False


def run_tesseract(image_path: Path, out_base: Path, lang: str, psm: str, output_format: str):
    cmd = [
        "tesseract",
        str(image_path),
        str(out_base),
        "-l",
        lang,
        "--psm",
        psm,
        "-c",
        "preserve_interword_spaces=1",
    ]

    if output_format == "txt":
        pass
    elif output_format == "md":
        # Tesseract writes plain text; caller renames to .md.
        pass
    elif output_format == "docx":
        # Caller handles docx assembly from plain text.
        pass
    elif output_format == "pdf":
        cmd.append("pdf")
    elif output_format == "hocr":
        cmd.append("hocr")
    else:
        raise ValueError(f"Unsupported format: {output_format}")

    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        raise RuntimeError(result.stderr.strip() or "Tesseract failed")


def text_to_docx(txt_path: Path, docx_path: Path, title: str):
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("python-docx is required for DOCX output. Install: pip install python-docx") from exc

    content = txt_path.read_text(encoding="utf-8", errors="replace")
    doc = Document()
    doc.add_heading(title, level=1)
    for line in content.splitlines():
        doc.add_paragraph(line)
    doc.save(docx_path)


class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Myanmar OCR Folder Tool (Tesseract)")
        self.geometry("860x560")

        self.input_dir = tk.StringVar(value=str(Path.cwd()))
        self.output_dir = tk.StringVar(value=str(Path.cwd() / "ocr_output"))
        self.lang = tk.StringVar(value="mya")
        self.psm = tk.StringVar(value="6")
        self.output_format = tk.StringVar(value="txt")
        self.combine_to_single = tk.BooleanVar(value=True)
        self.filename_filter = tk.StringVar(value="")
        self.page_range = tk.StringVar(value="")

        self._build_ui()

    def _build_ui(self):
        root = ttk.Frame(self, padding=12)
        root.pack(fill=tk.BOTH, expand=True)

        row = 0
        ttk.Label(root, text="Input folder").grid(row=row, column=0, sticky="w")
        ttk.Entry(root, textvariable=self.input_dir, width=85).grid(row=row, column=1, sticky="ew", padx=8)
        ttk.Button(root, text="Browse", command=self.pick_input_dir).grid(row=row, column=2, sticky="e")

        row += 1
        ttk.Label(root, text="Output folder").grid(row=row, column=0, sticky="w", pady=(8, 0))
        ttk.Entry(root, textvariable=self.output_dir, width=85).grid(
            row=row, column=1, sticky="ew", padx=8, pady=(8, 0)
        )
        ttk.Button(root, text="Browse", command=self.pick_output_dir).grid(row=row, column=2, sticky="e", pady=(8, 0))

        row += 1
        opts = ttk.Frame(root)
        opts.grid(row=row, column=0, columnspan=3, sticky="ew", pady=(12, 0))
        opts.columnconfigure(7, weight=1)

        ttk.Label(opts, text="Language").grid(row=0, column=0, sticky="w")
        ttk.Entry(opts, textvariable=self.lang, width=10).grid(row=0, column=1, sticky="w", padx=(6, 16))

        ttk.Label(opts, text="PSM").grid(row=0, column=2, sticky="w")
        ttk.Combobox(
            opts,
            textvariable=self.psm,
            values=["3", "4", "5", "6", "11", "12"],
            width=8,
            state="readonly",
        ).grid(row=0, column=3, sticky="w", padx=(6, 16))

        ttk.Label(opts, text="Output").grid(row=0, column=4, sticky="w")
        ttk.Combobox(
            opts,
            textvariable=self.output_format,
            values=["txt", "md", "docx", "pdf", "hocr"],
            width=10,
            state="readonly",
        ).grid(row=0, column=5, sticky="w", padx=(6, 16))

        ttk.Checkbutton(opts, text="Combine into a single file", variable=self.combine_to_single).grid(
            row=0, column=6, sticky="w"
        )

        row += 1
        filters = ttk.Frame(root)
        filters.grid(row=row, column=0, columnspan=3, sticky="ew", pady=(10, 0))
        filters.columnconfigure(5, weight=1)

        ttk.Label(filters, text="Filename contains").grid(row=0, column=0, sticky="w")
        ttk.Entry(filters, textvariable=self.filename_filter, width=24).grid(row=0, column=1, sticky="w", padx=(6, 16))

        ttk.Label(filters, text="Page range").grid(row=0, column=2, sticky="w")
        ttk.Entry(filters, textvariable=self.page_range, width=24).grid(row=0, column=3, sticky="w", padx=(6, 16))
        ttk.Label(filters, text="e.g. 3-10, 25, 40-").grid(row=0, column=4, sticky="w")

        row += 1
        ttk.Label(
            root,
            text="Tip: use 'mya' or 'mya+eng'. For layout fidelity, 'pdf' or 'hocr' is usually best.",
        ).grid(row=row, column=0, columnspan=3, sticky="w", pady=(8, 0))

        row += 1
        btns = ttk.Frame(root)
        btns.grid(row=row, column=0, columnspan=3, sticky="ew", pady=(12, 0))
        ttk.Button(btns, text="Start OCR", command=self.start_ocr, width=18).pack(side=tk.LEFT)
        ttk.Button(btns, text="Check Tesseract", command=self.check_tesseract, width=18).pack(side=tk.LEFT, padx=8)
        ttk.Button(btns, text="Myanmar Book Preset", command=self.apply_myanmar_preset, width=22).pack(
            side=tk.LEFT, padx=8
        )

        row += 1
        self.progress = ttk.Progressbar(root, mode="determinate")
        self.progress.grid(row=row, column=0, columnspan=3, sticky="ew", pady=(12, 0))

        row += 1
        self.log = tk.Text(root, height=20, wrap=tk.WORD)
        self.log.grid(row=row, column=0, columnspan=3, sticky="nsew", pady=(8, 0))
        root.columnconfigure(1, weight=1)
        root.rowconfigure(row, weight=1)

    def pick_input_dir(self):
        p = filedialog.askdirectory(initialdir=self.input_dir.get() or str(Path.cwd()))
        if p:
            self.input_dir.set(p)

    def pick_output_dir(self):
        p = filedialog.askdirectory(initialdir=self.output_dir.get() or str(Path.cwd()))
        if p:
            self.output_dir.set(p)

    def append_log(self, text: str):
        self.log.insert(tk.END, text + "\n")
        self.log.see(tk.END)
        self.update_idletasks()

    def check_tesseract(self):
        try:
            result = subprocess.run(["tesseract", "--list-langs"], capture_output=True, text=True)
            if result.returncode != 0:
                raise RuntimeError(result.stderr.strip() or "Tesseract not found.")
            self.append_log(result.stdout.strip())
            messagebox.showinfo("Tesseract", "Tesseract is available. Language list printed in log.")
        except FileNotFoundError:
            messagebox.showerror("Tesseract", "tesseract command not found in PATH.")
        except Exception as exc:
            messagebox.showerror("Tesseract", str(exc))

    def start_ocr(self):
        worker = threading.Thread(target=self.run_ocr, daemon=True)
        worker.start()

    def apply_myanmar_preset(self):
        self.lang.set("mya")
        self.psm.set("6")
        self.output_format.set("pdf")
        self.combine_to_single.set(False)
        self.append_log("Applied preset: lang=mya, psm=6, output=pdf, combine=off")

    def run_ocr(self):
        try:
            in_dir = Path(self.input_dir.get().strip())
            out_dir = Path(self.output_dir.get().strip())
            lang = self.lang.get().strip() or "mya"
            psm = self.psm.get().strip() or "6"
            fmt = self.output_format.get().strip()
            combine = self.combine_to_single.get()
            filename_filter = self.filename_filter.get().strip()
            page_range_raw = self.page_range.get().strip()

            if not in_dir.exists() or not in_dir.is_dir():
                messagebox.showerror("Input error", "Input folder does not exist.")
                return

            out_dir.mkdir(parents=True, exist_ok=True)
            images = list_images(in_dir)

            try:
                ranges = parse_page_range(page_range_raw)
            except ValueError as exc:
                raise RuntimeError("Invalid page range. Example: 3-10, 25, 40-") from exc
            if filename_filter:
                needle = filename_filter.lower()
                images = [img for img in images if needle in img.name.lower()]
            if ranges is not None:
                filtered = []
                for img in images:
                    num = extract_trailing_number(img.stem)
                    if num is None:
                        continue
                    if number_in_ranges(num, ranges):
                        filtered.append(img)
                images = filtered

            if not images:
                messagebox.showwarning("No images", "No image files found in input folder.")
                return

            self.log.delete("1.0", tk.END)
            self.append_log(f"Found {len(images)} image(s).")
            self.progress["maximum"] = len(images)
            self.progress["value"] = 0

            # For docx we first generate per-page txt then merge.
            if combine and fmt in {"txt", "md", "docx"}:
                combined_text_parts = []
            else:
                combined_text_parts = None

            for idx, image in enumerate(images, start=1):
                self.append_log(f"[{idx}/{len(images)}] OCR: {image.name}")
                base = out_dir / image.stem
                run_tesseract(image, base, lang, psm, "txt" if fmt == "docx" else fmt)

                if fmt == "md":
                    txt_path = base.with_suffix(".txt")
                    md_path = base.with_suffix(".md")
                    if txt_path.exists():
                        txt_path.rename(md_path)
                elif fmt == "docx":
                    txt_path = base.with_suffix(".txt")
                    docx_path = base.with_suffix(".docx")
                    if txt_path.exists() and not combine:
                        text_to_docx(txt_path, docx_path, image.name)
                        txt_path.unlink(missing_ok=True)

                if combined_text_parts is not None:
                    txt_path = base.with_suffix(".txt")
                    if txt_path.exists():
                        content = txt_path.read_text(encoding="utf-8", errors="replace")
                        combined_text_parts.append(f"===== {image.name} =====\n{content}\n")
                        if fmt == "docx":
                            txt_path.unlink(missing_ok=True)

                self.progress["value"] = idx
                self.update_idletasks()

            if combined_text_parts is not None:
                self._write_combined_file(out_dir, in_dir.name, fmt, combined_text_parts)

            self.append_log("Done.")
            messagebox.showinfo("OCR finished", f"OCR completed.\nOutput folder:\n{out_dir}")

        except FileNotFoundError:
            messagebox.showerror("Tesseract", "tesseract command not found in PATH.")
        except Exception as exc:
            self.append_log(f"ERROR: {exc}")
            messagebox.showerror("OCR failed", str(exc))

    def _write_combined_file(self, out_dir: Path, folder_name: str, fmt: str, parts):
        if fmt == "txt":
            out_file = out_dir / "ocr_combined.txt"
            out_file.write_text("\n".join(parts), encoding="utf-8")
            self.append_log(f"Combined file: {out_file.name}")
            return

        if fmt == "md":
            out_file = out_dir / "ocr_combined.md"
            out_file.write_text("\n".join(parts), encoding="utf-8")
            self.append_log(f"Combined file: {out_file.name}")
            return

        if fmt == "docx":
            try:
                from docx import Document
            except Exception as exc:
                raise RuntimeError(
                    "python-docx is required for DOCX output. Install: pip install python-docx"
                ) from exc

            doc = Document()
            doc.add_heading(folder_name, level=1)
            for block in parts:
                lines = block.splitlines()
                if not lines:
                    continue
                doc.add_heading(lines[0], level=2)
                for line in lines[1:]:
                    doc.add_paragraph(line)
            out_file = out_dir / f"{folder_name}.docx"
            doc.save(out_file)
            self.append_log(f"Combined file: {out_file.name}")
            return


if __name__ == "__main__":
    App().mainloop()
