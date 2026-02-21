import argparse
import re
from copy import deepcopy
from pathlib import Path

from docx import Document


def natural_key(text: str):
    return [int(part) if part.isdigit() else part.lower() for part in re.split(r"(\d+)", text)]


def remove_default_blank_paragraph(doc: Document):
    if len(doc.paragraphs) == 1 and not doc.paragraphs[0].text and not doc.tables:
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)


def append_document(master: Document, src: Document, add_page_break: bool):
    if add_page_break:
        master.add_page_break()

    for element in src.element.body:
        # Skip section properties from source docs to avoid repeated section definitions.
        if element.tag.endswith("}sectPr"):
            continue
        master.element.body.append(deepcopy(element))


def combine_docx(input_dir: Path, output_file: Path):
    docs = sorted(
        [
            p
            for p in input_dir.glob("*.docx")
            if not p.name.startswith("~$") and p.resolve() != output_file.resolve()
        ],
        key=lambda p: natural_key(p.name),
    )
    if not docs:
        raise RuntimeError("No .docx files found in input folder.")

    master = Document()
    remove_default_blank_paragraph(master)
    master.add_heading("Master", level=1)

    first = True
    for path in docs:
        src = Document(path)
        master.add_heading(path.stem, level=2)
        append_document(master, src, add_page_break=not first)
        first = False

    output_file.parent.mkdir(parents=True, exist_ok=True)
    master.save(output_file)
    return len(docs)


def main():
    parser = argparse.ArgumentParser(description="Combine multiple DOCX files into one Master DOCX.")
    parser.add_argument(
        "input_dir",
        nargs="?",
        default=".",
        help="Folder containing DOCX files (default: current folder)",
    )
    parser.add_argument(
        "--output",
        default=None,
        help="Output DOCX path (default: <input_folder_name>_Master.docx in input folder)",
    )
    args = parser.parse_args()

    input_dir = Path(args.input_dir).resolve()
    if not input_dir.exists() or not input_dir.is_dir():
        raise RuntimeError(f"Input folder not found: {input_dir}")

    default_output = input_dir / f"{input_dir.name}_Master.docx"
    output_file = Path(args.output).resolve() if args.output else default_output

    count = combine_docx(input_dir, output_file)
    print(f"Combined {count} files into: {output_file}")


if __name__ == "__main__":
    main()
